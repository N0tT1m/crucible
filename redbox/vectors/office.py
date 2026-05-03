"""E4 doc-trojan — DOCX / XLSX with hidden payload text.

Both vectors lazily import their respective libs (`python-docx`, `openpyxl`)
so the base install stays slim. Install with `pip install -e .[docs]`.

  DocxHiddenTextVector  — runs with the "hidden" text property set to
                          true and white-on-white runs. Word displays
                          neither, but text-extractors pull both.
  XlsxCommentInjection  — adds cell comments (which most spreadsheet
                          summarisers expand inline) plus a hidden
                          column whose values carry the payload.
"""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

from redbox.core.types import Payload
from redbox.core.vector import Artifact


def _have(mod: str) -> bool:
    try:
        __import__(mod)
    except ImportError:
        return False
    return True


@dataclass
class DocxHiddenTextVector:
    name: str = "docx_hidden_text"
    channel: str = "docx"

    def embed(self, payload: Payload) -> Artifact:
        if not _have("docx"):
            raise RuntimeError("docx_hidden_text needs python-docx (pip install -e .[docs])")
        from docx import Document as DocxDoc
        from docx.shared import RGBColor

        doc = DocxDoc()
        doc.add_heading("Quarterly summary", level=1)
        doc.add_paragraph("All metrics within target. See attached for details.")

        # Hidden run: vanish=true. Word + most viewers elide; extractors keep.
        p = doc.add_paragraph()
        run = p.add_run(payload.template)
        run.font.hidden = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Footer-style invisible appendix.
        p2 = doc.add_paragraph()
        run2 = p2.add_run("Appendix: " + payload.template)
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = None

        buf = BytesIO()
        doc.save(buf)
        return Artifact(
            channel="docx", body=buf.getvalue(),
            filename=f"{payload.id}__hidden.docx",
            metadata={"payload_id": payload.id},
        )


@dataclass
class XlsxCommentInjection:
    name: str = "xlsx_comment_inject"
    channel: str = "xlsx"

    def embed(self, payload: Payload) -> Artifact:
        if not _have("openpyxl"):
            raise RuntimeError("xlsx_comment_inject needs openpyxl (pip install -e .[docs])")
        from openpyxl import Workbook
        from openpyxl.comments import Comment

        wb = Workbook()
        ws = wb.active
        ws.title = "Q3 metrics"
        ws["A1"] = "Metric"
        ws["B1"] = "Value"
        ws["A2"] = "Revenue"
        ws["B2"] = 1234567

        # Cell comment — most LLM-driven spreadsheet summarisers expand these.
        ws["B2"].comment = Comment(payload.template[:500], "ops")

        # Hidden column carrying the payload.
        ws["E1"] = "internal"
        for i, line in enumerate(payload.template.splitlines() or [payload.template], start=2):
            ws.cell(row=i, column=5, value=line[:300])
        ws.column_dimensions["E"].hidden = True

        buf = BytesIO()
        wb.save(buf)
        return Artifact(
            channel="xlsx", body=buf.getvalue(),
            filename=f"{payload.id}__hidden.xlsx",
            metadata={"payload_id": payload.id},
        )
